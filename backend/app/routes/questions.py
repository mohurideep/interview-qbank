import uuid
from io import BytesIO
from datetime import datetime, timedelta

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document

from ..db import get_db
from ..settings import settings
from .. import crud
from ..schemas import QuestionCreate, QuestionUpdate, QuestionOut

router = APIRouter(prefix="/v1/questions", tags=["questions"])


def _user_id() -> uuid.UUID:
    """
    Single-user mode: all data belongs to DEFAULT_USER_ID.
    Set in .env:
      DEFAULT_USER_ID=00000000-0000-0000-0000-000000000001
    """
    return uuid.UUID(settings.DEFAULT_USER_ID)


def _to_out(q) -> QuestionOut:
    history = [event.studied_at for event in (q.study_events or [])]
    if not history and q.studied_at is not None:
        history = [q.studied_at]
    return QuestionOut(
        id=q.id,
        parent_id=q.parent_id,
        studied_at=q.studied_at,
        studied_count=len(history),
        studied_history=history,
        question_text=q.question_text,
        answer_md=q.answer_md,
        difficulty=q.difficulty,
        source=q.source,
        is_flagged=q.is_flagged,
        tags=[t.name for t in q.tags],
        created_at=q.created_at,
        updated_at=q.updated_at,
        review_count=q.review_count,
        mastery_score=q.mastery_score,
        next_review_at=q.next_review_at,
    )


def _question_depth(items: list) -> dict[uuid.UUID, int]:
    by_id = {q.id: q for q in items}
    memo: dict[uuid.UUID, int] = {}
    visiting: set[uuid.UUID] = set()

    def depth(qid: uuid.UUID) -> int:
        if qid in memo:
            return memo[qid]
        if qid in visiting:
            return 0

        visiting.add(qid)
        q = by_id[qid]
        if q.parent_id and q.parent_id in by_id:
            val = depth(q.parent_id) + 1
        else:
            val = 0
        visiting.remove(qid)
        memo[qid] = val
        return val

    for q in items:
        depth(q.id)

    return memo


def _build_docx(items: list, title: str) -> BytesIO:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"Generated at {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    if not items:
        document.add_paragraph("No questions found.")
    else:
        depth_map = _question_depth(items)
        for idx, q in enumerate(items, start=1):
            indent = "  " * depth_map.get(q.id, 0)
            q_line = document.add_paragraph()
            q_prefix = q_line.add_run(f"Q{idx}: ")
            q_prefix.bold = True
            q_line.add_run(f"{indent}{q.question_text}")

            if q.source:
                document.add_paragraph(f"Source: {q.source}")
            if q.tags:
                document.add_paragraph(f"Tags: {', '.join(t.name for t in q.tags)}")

            a_line = document.add_paragraph()
            a_prefix = a_line.add_run("A: ")
            a_prefix.bold = True
            a_line.add_run(q.answer_md or "No answer provided.")
            document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _ensure_valid_parent(
    db: Session,
    user_id: uuid.UUID,
    parent_id: uuid.UUID | None,
    child_id: uuid.UUID | None = None,
) -> None:
    if parent_id is None:
        return

    if child_id is not None and parent_id == child_id:
        raise HTTPException(status_code=400, detail="A question cannot be its own parent")

    parent = crud.get_question(db, user_id, parent_id)
    if not parent:
        raise HTTPException(status_code=404, detail="Parent question not found")

    if child_id is None:
        return

    ancestor = parent
    while ancestor.parent_id is not None:
        if ancestor.parent_id == child_id:
            raise HTTPException(status_code=400, detail="Parent relationship would create a cycle")
        ancestor = crud.get_question(db, user_id, ancestor.parent_id)
        if not ancestor:
            break


@router.post("", response_model=QuestionOut)
def create(payload: QuestionCreate, db: Session = Depends(get_db)):
    user_id = _user_id()
    _ensure_valid_parent(db, user_id, payload.parent_id)
    try:
        q = crud.create_question(db, user_id, payload)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return _to_out(q)


@router.get("", response_model=list[QuestionOut])
def list_(
    db: Session = Depends(get_db),
    search: str | None = Query(default=None),
    tag: str | None = Query(default=None),
    tags: str | None = Query(default=None),
    source: str | None = Query(default=None),
    flagged: bool | None = Query(default=None),
    due_only: bool | None = Query(default=False),
):
    tag_filter = tags if tags is not None else tag
    items = crud.list_questions(db, _user_id(), search, tag_filter, source, flagged)

    if due_only:
        now = datetime.utcnow()
        items = [q for q in items if (q.next_review_at is None) or (q.next_review_at <= now)]

    return [_to_out(q) for q in items]


@router.get("/export")
def export_questions(
    db: Session = Depends(get_db),
    thread_id: uuid.UUID | None = Query(default=None),
    search: str | None = Query(default=None),
    tag: str | None = Query(default=None),
    tags: str | None = Query(default=None),
    source: str | None = Query(default=None),
    flagged: bool | None = Query(default=None),
    due_only: bool | None = Query(default=False),
):
    user_id = _user_id()
    if thread_id is not None:
        items = crud.list_thread_questions(db, user_id, thread_id)
        if not items:
            raise HTTPException(status_code=404, detail="Thread not found")
        title = f"Interview QBank Thread Export ({thread_id})"
        filename = f"thread-{thread_id}.docx"
    else:
        tag_filter = tags if tags is not None else tag
        items = crud.list_questions(db, user_id, search, tag_filter, source, flagged)
        if due_only:
            now = datetime.utcnow()
            items = [q for q in items if (q.next_review_at is None) or (q.next_review_at <= now)]
        title = "Interview QBank Export"
        filename = f"interview-qbank-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}.docx"

    stream = _build_docx(items, title=title)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/suggestions", response_model=list[str])
def suggestions(
    db: Session = Depends(get_db),
    field: str = Query(..., description='One of: "source", "tag"'),
    q: str = Query(default=""),
    limit: int = Query(default=8, ge=1, le=30),
):
    field = field.strip().lower()
    if field == "source":
        return crud.list_source_suggestions(db, _user_id(), q, limit)
    if field == "tag":
        return crud.list_tag_suggestions(db, _user_id(), q, limit)
    raise HTTPException(status_code=400, detail='Invalid field. Use "source" or "tag".')


@router.get("/{qid}", response_model=QuestionOut)
def get_one(qid: uuid.UUID, db: Session = Depends(get_db)):
    q = crud.get_question(db, _user_id(), qid)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    return _to_out(q)


@router.patch("/{qid}", response_model=QuestionOut)
def patch(qid: uuid.UUID, payload: QuestionUpdate, db: Session = Depends(get_db)):
    user_id = _user_id()
    q = crud.get_question(db, user_id, qid)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")

    if "parent_id" in payload.model_fields_set:
        _ensure_valid_parent(db, user_id, payload.parent_id, child_id=qid)

    try:
        q = crud.update_question(db, q, user_id, payload)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return _to_out(q)


@router.delete("/{qid}")
def delete(qid: uuid.UUID, db: Session = Depends(get_db)):
    q = crud.get_question(db, _user_id(), qid)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    db.delete(q)
    db.commit()
    return {"status": "deleted"}


@router.post("/{qid}/review")
def review(
    qid: uuid.UUID,
    rating: str = Query(..., description='One of: "forgot", "almost", "knew"'),
    db: Session = Depends(get_db),
):
    q = crud.get_question(db, _user_id(), qid)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")

    q.review_count = (q.review_count or 0) + 1

    rating = rating.lower().strip()
    if rating == "forgot":
        q.mastery_score = float(q.mastery_score or 0.0) - 0.3
        interval_days = 1
    elif rating == "almost":
        q.mastery_score = float(q.mastery_score or 0.0) + 0.1
        interval_days = 3
    elif rating == "knew":
        q.mastery_score = float(q.mastery_score or 0.0) + 0.3
        interval_days = 7
    else:
        raise HTTPException(status_code=400, detail='Invalid rating. Use "forgot", "almost", or "knew".')

    q.mastery_score = max(0.0, min(5.0, q.mastery_score))
    q.next_review_at = datetime.utcnow() + timedelta(days=interval_days)
    q.updated_at = datetime.utcnow()

    db.commit()
    return {"status": "ok", "next_review_at": q.next_review_at, "mastery_score": q.mastery_score}
